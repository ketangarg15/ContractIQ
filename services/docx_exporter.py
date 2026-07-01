"""
Word (.docx) report exporter service.
"""

import io
import json
from docx import Document
from docx.shared import Inches, Pt, RGBColor

def _parse_json(val):
    if isinstance(val, str):
        try:
            return json.loads(val)
        except Exception:
            return val
    return val

def generate_docx_report(
    filename: str,
    summary: dict | str | None,
    clauses: dict | str | None,
    risk_analysis: dict | str | None,
    risk_score: float | None,
    contract_type: str | None,
    obligations: dict | str | None,
    entities: dict | str | None,
    red_flags: dict | str | None,
    compliance: dict | str | None,
    negotiations: dict | str | None,
) -> bytes:
    """Generates a Microsoft Word document containing contract analytics and negotiation redlines."""
    summary = _parse_json(summary)
    clauses = _parse_json(clauses)
    risk_analysis = _parse_json(risk_analysis)
    obligations = _parse_json(obligations)
    entities = _parse_json(entities)
    red_flags = _parse_json(red_flags)
    compliance = _parse_json(compliance)
    negotiations = _parse_json(negotiations)

    doc = Document()
    
    # Title Section
    title = doc.add_heading('ContractIQ — Legal Analysis & Redline Report', level=0)
    title.style.font.size = Pt(24)
    title.style.font.color.rgb = RGBColor(37, 99, 235)  # Premium blue color
    
    # Metadata Block
    doc.add_heading('1. Metadata & Classification', level=1)
    doc.add_paragraph(f"Contract Filename: {filename}")
    doc.add_paragraph(f"Classified Contract Type: {contract_type or 'General Agreement'}")
    doc.add_paragraph(f"Overall Risk Score: {int(risk_score) if risk_score is not None else 'N/A'} / 100")
    
    # Executive Summary
    doc.add_heading('2. Executive Summary', level=1)
    if summary and isinstance(summary, dict) and "narrative" in summary:
        doc.add_paragraph(summary["narrative"])
    else:
        doc.add_paragraph("No narrative summary available.")
        
    # Key Obligations
    doc.add_heading('3. Key Obligations & Deadlines', level=1)
    if obligations and isinstance(obligations, dict) and obligations.get("obligations"):
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Light Shading Accent 1'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Type'
        hdr_cells[1].text = 'Description'
        hdr_cells[2].text = 'Responsible Party'
        hdr_cells[3].text = 'Due Date'
        
        for obs in obligations.get("obligations", []):
            row_cells = table.add_row().cells
            row_cells[0].text = str(obs.get("obligation_type", ""))
            row_cells[1].text = str(obs.get("description", ""))
            row_cells[2].text = str(obs.get("responsible_party", ""))
            row_cells[3].text = str(obs.get("due_date") or "N/A")
    else:
        doc.add_paragraph("No specific obligations listed.")
        
    # Negotiation Redlines Annex
    doc.add_heading('4. AI Negotiation Suggestions & Redlines', level=1)
    if negotiations and isinstance(negotiations, dict) and negotiations.get("suggestions"):
        for idx, sug in enumerate(negotiations.get("suggestions", []), 1):
            doc.add_heading(f"Deviation #{idx}: {sug.get('clause_name', 'Clause')}", level=2)
            doc.add_paragraph(f"Reference: {sug.get('reference') or 'N/A'}")
            doc.add_paragraph(f"Priority: {sug.get('priority', 'Medium')} | Negotiability: {sug.get('negotiability', 'Moderate')}")
            
            p_curr = doc.add_paragraph()
            p_curr.add_run("Current Wording: ").bold = True
            p_curr.add_run(sug.get('current_wording', ''))
            
            p_sug = doc.add_paragraph()
            p_sug.add_run("Suggested Redline: ").bold = True
            p_sug.add_run(sug.get('suggested_wording', ''))
            
            p_rat = doc.add_paragraph()
            p_rat.add_run("Rationale: ").bold = True
            p_rat.add_run(sug.get('rationale', ''))
    else:
        doc.add_paragraph("No negotiation redline recommendations available.")
        
    # Save document to bytes buffer
    file_stream = io.BytesIO()
    doc.save(file_stream)
    return file_stream.getvalue()
